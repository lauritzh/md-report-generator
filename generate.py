#!/usr/bin/env python3
#
# A simple, self-contained report generator for penetration testing reports.
#
# Lauritz Holtmann, (c) 2022 - 2025
#

import io
import os
import re
import sys
import yaml
import pdfkit
import argparse
import datetime
import markdown
import xlsxwriter
import matplotlib.pyplot as plt
from cvss import CVSS4
from datetime import date
from string import Template
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Constants
content_dir = "content/"
findings_dir = "findings/"
output_dir = "output/"
boilerplate_dir = "boilerplate/"
page_break = '\n\n<div style = "display:block; clear:both; page-break-after:always;"></div>\n\n'
cvss_version = "4.0"
cvss_base_metrics = ["AV", "AC", "AT", "PR", "UI", "VC", "VI", "VA", "SC", "SI", "SA"]

# Global Variables
config = {}
findings = []
report_md = ""
report_html = ""
findings_list = ""
cover_location = ""
generated_piechart = ""
total_findings = critical_findings = high_findings = medium_findings = low_findings = none_findings = 0

# Set Base-URL to current working directory 
# Makes including images to report more easy by simply referencing images/test.png
base_href = "file://{}/".format(os.getcwd())

def wrap_html_document(body_html, base_href=None):
	"""Wrap bare HTML fragments with a minimal document structure and base tag."""
	base_tag = '<base href="{}">'.format(base_href) if base_href else ''
	return """<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
{}
</head>
<body>
{}
</body>
</html>""".format(base_tag, body_html)

def resolve_finding_id(finding, counter):
	"""Return canonical identifier for a finding, honoring metadata overrides."""
	if finding.get("finding_id"):
		override = str(finding["finding_id"]).strip()
		if override:
			return override
	return "PEN{}{:04d}".format(date.today().year, counter+1)

def display_finding_id(finding, counter):
	"""Return a display-friendly finding id (prefixed with # if needed)."""
	finding_id = resolve_finding_id(finding, counter)
	return finding_id if finding_id.startswith("#") else "#{}".format(finding_id)

def build_cvss4_vector(cvss_config):
	"""Create a normalized CVSS 4.0 vector from either a dict of metrics or a raw vector."""
	if not isinstance(cvss_config, dict):
		raise ValueError("CVSS configuration must be a mapping")

	vector_value = cvss_config.get("vector")
	if vector_value:
		vector = str(vector_value).strip()
		if not vector:
			raise ValueError("Provided CVSS vector is empty")
		vector_body = vector
		if vector_body.startswith("CVSS:"):
			vector_body = vector_body.split("/", 1)[1] if "/" in vector_body else ""
		vector_body = vector_body.lstrip("/")
		if not vector_body:
			raise ValueError("Provided CVSS vector is incomplete")
		return "CVSS:{}/{}".format(cvss_version, vector_body)

	missing = [metric for metric in cvss_base_metrics if metric not in cvss_config]
	if missing:
		raise ValueError("Missing CVSS metrics: {}".format(", ".join(missing)))

	def format_metric(metric):
		value = str(cvss_config[metric]).strip().upper()
		if not value:
			raise ValueError("Empty value for metric {}".format(metric))
		return "{}:{}".format(metric, value)

	return "CVSS:{}/{}".format(cvss_version, "/".join(format_metric(m) for m in cvss_base_metrics))

def markdown_links_to_text(text):
	"""Convert Markdown inline links to plain text with URL in parentheses."""
	if not text:
		return ""
	link_pattern = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
	return link_pattern.sub(lambda m: "{} ({})".format(m.group(1).strip(), m.group(2).strip()), text)

def ensure_code_block_style(document):
	"""Ensure a dedicated code-block style exists and return it."""
	style_name = "CodeBlock"
	try:
		return document.styles[style_name]
	except KeyError:
		style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
		style.font.name = "Courier New"
		style.font.size = Pt(10)
		style.paragraph_format.left_indent = Pt(12)
		style.paragraph_format.space_before = Pt(6)
		style.paragraph_format.space_after = Pt(6)
		return style

def apply_code_block_shading(paragraph, fill="F2F2F2"):
	"""Apply a soft gray background to a paragraph to mimic code block styling."""
	ppr = paragraph._p.get_or_add_pPr()
	shading = OxmlElement('w:shd')
	shading.set(qn('w:val'), 'clear')
	shading.set(qn('w:color'), 'auto')
	shading.set(qn('w:fill'), fill)
	ppr.append(shading)

def add_code_block_paragraph(document, lines):
	"""Insert a formatted code block paragraph into the document."""
	if not lines:
		return
	style = ensure_code_block_style(document)
	paragraph = document.add_paragraph(style=style)
	paragraph.alignment = None
	run = paragraph.add_run("\n".join(lines))
	run.font.name = "Courier New"
	run.font.size = Pt(10)
	apply_code_block_shading(paragraph)

def add_markdown_content_to_doc(document, markdown_text):
	"""Render a limited subset of Markdown to the Word document."""
	if not markdown_text:
		return

	lines = markdown_text.strip().splitlines()
	in_code_block = False
	code_lines = []

	for raw_line in lines:
		line = raw_line.rstrip("\n")
		stripped = line.strip()

		if stripped.startswith("```"):
			if in_code_block:
				add_code_block_paragraph(document, code_lines)
				code_lines = []
				in_code_block = False
			else:
				in_code_block = True
			continue

		if in_code_block:
			code_lines.append(line)
			continue

		if not stripped:
			document.add_paragraph("")
			continue

		if stripped.startswith("#"):
			level = len(stripped) - len(stripped.lstrip("#"))
			heading_text = stripped[level:].strip()
			if heading_text:
				heading_level = max(1, min(level, 4))
				document.add_heading(markdown_links_to_text(heading_text), level=heading_level)
			continue

		bullet_prefix = None
		for prefix in ("* ", "- ", "• ", "● "):
			if stripped.startswith(prefix):
				bullet_prefix = prefix
				break

		if bullet_prefix is not None:
			bullet_text = stripped[len(bullet_prefix):].strip()
			document.add_paragraph(markdown_links_to_text(bullet_text), style='List Bullet')
			continue

		document.add_paragraph(markdown_links_to_text(line))

	if in_code_block and code_lines:
		add_code_block_paragraph(document, code_lines)

def init():
	"""Initialize the report generator, load config from config.yaml"""
	global config
	# Parse Config
	with open('config.yaml') as f:
		config = yaml.load(f, Loader=yaml.FullLoader)
		print("Config options: {}".format(config))
		f.close()

def generate_report():
	"""Generate the PDF report"""
	global report_html
	# Generate Markdown Report
	generate_markdown_report()
	# Generate PDF Report
	generate_pdf_report(report_html)

def generate_markdown_report():
	"""Generate the Markdown report from the Markdown template and findings"""
	global config, content_dir, cover_location, findings, findings_dir, report_md, report_html, total_findings, critical_findings, high_findings, medium_findings, low_findings, none_findings, base_href

	# Glue: Collect files and build Markdown report
	with open(content_dir + 'introduction.md') as f:
		report_md += f.read()
		report_md += page_break
		f.close()

	with open(content_dir + 'history.md') as f:
		report_md += f.read()
		report_md += page_break
		f.close()

	with open(content_dir + 'scope.md') as f:
		report_md += f.read()
		report_md += page_break
		f.close()

	with open(content_dir + 'technical-details.md') as f:
		report_md += f.read()
		report_md += page_break
		f.close()

	# Insert Placeholders
	report_md = report_md.format(
		title = config["title"], 
		author = config["author"], 
		customer = config["customer"],
		vendor = config["vendor"],
		critical_findings = "{critical_findings}",
		high_findings = "{high_findings}",
		medium_findings = "{medium_findings}",
		low_findings = "{low_findings}",
		piechart = "{piechart}",
		findings_list = "{findings_list}",
		fixed_findings_list = "{fixed_findings_list}"
	)

	# Process Findings
	process_findings()

	# Determine Statistics and Render Pie Chart
	print("Generating Pie Chart...")

	## Data for the pie chart
	# Optional: Include informational findnigs
	#labels = ['Critical', 'High', 'Medium', 'Low', 'None']
	#sizes = [critical_findings, high_findings, medium_findings, low_findings, none_findings]
	#colors = ['violet', 'red', 'orange', 'yellow', 'green']
	labels = ['Critical', 'High', 'Medium', 'Low']
	sizes = [critical_findings, high_findings, medium_findings, low_findings]
	colors = ['violet', 'red', 'orange', 'yellow']

	## Set font size and padding for legend
	plt.rcParams['font.size'] = 12
	plt.rcParams['legend.fontsize'] = 12
	plt.rcParams['font.family'] = 'sans-serif'
	plt.rcParams['font.sans-serif'] = ['avenir']
	plt.rcParams["figure.autolayout"] = True

	## Create the pie chart as an SVG in memory
	fig, ax = plt.subplots()
	ax.pie(sizes, labels=None, colors=colors, autopct=lambda pct: f"{pct:.1f}%" if pct > 0 else '')
	ax.axis('equal')
	### Set legend
	plt.subplots_adjust(left=0.1, right=0.5)
	ax.legend(labels, loc='center left', bbox_to_anchor=(1, 0.5), title='Distribution of Findings by Severity')
	leg = ax.get_legend()
	leg._legend_box.align = "left"
	svg_io = io.BytesIO()
	plt.savefig(svg_io, format='svg')
	svg_io.seek(0)
	generated_piechart = svg_io.getvalue().decode('utf-8')

	## Create the detailed table of findings (exclude fixed)
	active_findings = [f for f in findings if not f.get("fixed", False)]
	fixed_findings = [f for f in findings if f.get("fixed", False)]

	generated_table_of_findings = ""
	for counter,finding in enumerate(active_findings):
		finding_id = finding.get("_display_id", display_finding_id(finding, counter))
		# Fill Template
		generated_table_of_findings += "* <b style='display:inline-block;width:100px'>{}</b> {}:\t{} ([CWE-{}](https://cwe.mitre.org/data/definitions/{}.html))\n".format(finding["cvss_severity"], finding_id, finding["title"], finding["CWE-ID"], finding["CWE-ID"])

	## Create list of fixed findings (titles only with IDs)
	generated_fixed_findings = ""
	for counter,finding in enumerate(fixed_findings):
		finding_id = finding.get("_display_id", display_finding_id(finding, counter))
		generated_fixed_findings += "* <b>Fixed</b> {}: {} ([CWE-{}](https://cwe.mitre.org/data/definitions/{}.html))\n".format(finding_id, finding["title"], finding["CWE-ID"], finding["CWE-ID"])

	# Insert Placeholders
	report_md = report_md.format(
		critical_findings = critical_findings,
		high_findings = high_findings,
		medium_findings = medium_findings,
		low_findings = low_findings,
		piechart = "{piechart}",
		findings_list = generated_table_of_findings,
		fixed_findings_list = generated_fixed_findings if generated_fixed_findings != "" else "_No fixed findings._"
	)

	# Append processed findings (including fixed) to report
	for counter,finding in enumerate(findings):
		print("Appending finding {}...".format(finding["title"]))
		# Fill Template
		finding_id = finding.get("_display_id", display_finding_id(finding, counter))
		report_md += finding_markdown(finding, finding_id)

	# Append Conclusion and Appendix
	with open(content_dir + 'conclusion.md') as f:
		report_md += f.read() 
		report_md += page_break
		f.close()

	with open(content_dir + 'appendix.md') as f:
		report_md += f.read() 
		report_md += page_break
		f.close()

	# Render Markdown: Convert to main report to HTML
	print("Render Markdown to HTML...")
	report_html = markdown.markdown(report_md, extensions=['fenced_code', 'codehilite', 'tables'])

	cover_location = "temp/cover_processed.html"
	with open(boilerplate_dir + 'cover.html') as f:
		cover_processed = Template(f.read()).safe_substitute(title=config["title"], author=config["author"], vendor=config["vendor"], date=datetime.datetime.now().strftime("%Y-%m-%d"), customer=config["customer"])
		f.close()

	with open(cover_location, 'w') as f:
		f.write(cover_processed)
		f.close()

	# Insert inlined SVG
	report_html = report_html.replace("{piechart}", generated_piechart)
	report_html = wrap_html_document(report_html, base_href)

def finding_markdown(finding, finding_id = "TBD"):
	"""Generate Markdown for a single finding"""
	is_fixed = finding.get("fixed", False)
	status_label = "Fixed" if is_fixed else "Open"
	title_suffix = " (Fixed)" if is_fixed else ""
	temp = """
### {}: {}{}

---

| Asset         | CWE                                                      | Status | Severity | CVSS v{cvss_version} Vector                                                                             |
|---------------|----------------------------------------------------------|--------|---------------------------------|----------------------------------------------------------------------------------------------|
| {} | [{}]({}) | **{}** | {} ({})                      | *{}* |

---

{}

	""".format(
		finding_id,
		finding["title"],
		title_suffix,
		finding["asset"],
		finding["CWE-ID"],
		finding["CWE-Link"],
		status_label,
		finding["cvss_severity"],
		finding["cvss_score"],
		finding["cvss_vector"],
		finding["description"],
		cvss_version=cvss_version
	)
	return temp + page_break

def process_findings():
	"""Process all findings and generate statistics"""
	global config, findings, findings_dir, total_findings, critical_findings, high_findings, medium_findings, low_findings, none_findings

	# Iterate over finding MD files, preprocess
	for file in os.listdir(findings_dir):
		if file.endswith(".md"):
			filename = os.fsdecode(file)
			with open(findings_dir + filename) as f:
				print("Processing finding {}...".format(filename))
				finding = {}

				# Map finding description from MD file
				finding["description"] = f.read()
				f.close()

				# Parse Properties from Header Section
				re_search = re.search(r"<!--[\r\n]([\s\S]*)[\r\n]-->", finding["description"])
				properties_yaml = re_search.group(1)
				properties = yaml.load(properties_yaml, Loader=yaml.FullLoader)
				# Cleanup: Remove properties
				finding["description"] = finding["description"].replace(re_search.group(0), "")

				# Map Properties
				finding["title"] = properties["title"]
				finding["asset"] = properties["asset"]
				finding["CWE-ID"] = properties["CWE-ID"]
				finding["CWE-Link"] = properties["CWE-Link"]
				if "finding_id" in properties:
					finding["finding_id"] = properties["finding_id"]
				# Optional fixed flag
				finding["fixed"] = bool(properties.get("fixed", False))

				# calculate CVSS score and severity
				try:
					cvss_vector = build_cvss4_vector(properties["cvss"])
				except Exception as exc:
					raise ValueError("Invalid CVSS configuration in {}: {}".format(filename, exc)) from exc

				c = CVSS4(cvss_vector)
				finding["cvss_vector"] = c.clean_vector()
				finding["cvss_score"] = c.scores()[0]
				finding["cvss_severity"] = c.severities()[0]

				findings.append(finding)
		else:
			print("File {} does not have correct file type .md".format(file))

	# Sort findings, CVSS Score descending
	def useScore(elem):
		return elem["cvss_score"]
	findings.sort(key=useScore,reverse=True)

	# Precompute resolved/display IDs for consistent reuse
	for idx, finding in enumerate(findings):
		resolved = resolve_finding_id(finding, idx)
		display_id = resolved if resolved.startswith("#") else "#{}".format(resolved)
		finding["_resolved_id"] = resolved
		finding["_display_id"] = display_id

	# Stats exclude fixed findings
	active_findings = [f for f in findings if not f.get("fixed", False)]
	total_findings = len(active_findings)
	critical_findings = len([finding for finding in active_findings if finding["cvss_severity"] == "Critical"])
	high_findings = len([finding for finding in active_findings if finding["cvss_severity"] == "High"])
	medium_findings = len([finding for finding in active_findings if finding["cvss_severity"] == "Medium"])
	low_findings = len([finding for finding in active_findings if finding["cvss_severity"] == "Low"])
	none_findings = len([finding for finding in active_findings if finding["cvss_severity"] == "None"])


def generate_excel_report():
	"""Generate Excel Report"""
	global config, findings, output_dir
	# Write findings to Excel file
	print("Generating Excel file...")
	excel_report = xlsxwriter.Workbook(output_dir + 'report.xlsx')
	excel_report_sheet = excel_report.add_worksheet("Findings")
	bold = excel_report.add_format({'bold': True})
	table_header = excel_report.add_format({'bold': True, 'bg_color': '#c8c8cf'})

	# Formats for Fixed column
	fixed_yes_fmt = excel_report.add_format({'bold': True, 'font_color': 'white', 'bg_color': '#2e7d32', 'align': 'center'})  # green
	fixed_no_fmt = excel_report.add_format({'bold': True, 'font_color': 'white', 'bg_color': '#c62828', 'align': 'center'})   # red

	# Formats for Severity column
	sev_crit_fmt = excel_report.add_format({'bold': True, 'font_color': 'white', 'bg_color': '#7b1fa2'})  # violet
	sev_high_fmt = excel_report.add_format({'bold': True, 'font_color': 'white', 'bg_color': '#d32f2f'})  # red
	sev_med_fmt  = excel_report.add_format({'bold': True, 'font_color': 'black', 'bg_color': '#f9a825'})  # orange/yellow
	sev_low_fmt  = excel_report.add_format({'bold': False, 'font_color': 'black', 'bg_color': '#fff59d'})  # light yellow
	sev_none_fmt = excel_report.add_format({'bold': False, 'font_color': 'black', 'bg_color': '#c8e6c9'})  # light green

	# Title
	excel_report_sheet.write(0, 0, "Pentest Report: {}".format(config["title"]), bold)
	excel_report_sheet.write(1, 0, "Author: {}".format(config["author"]))
	excel_report_sheet.write(2, 0, "Date: {}".format(datetime.datetime.now().strftime("%Y-%m-%d")))

	# Table Header
	excel_report_sheet.write(4, 0, "Finding-ID", table_header)
	excel_report_sheet.write(4, 1, "Fixed", table_header)
	excel_report_sheet.write(4, 2, "Severity", table_header)
	excel_report_sheet.write(4, 3, "Asset", table_header)
	excel_report_sheet.write(4, 4, "Title", table_header)

	# Column widths
	excel_report_sheet.set_column(0, 0, 16)  # Finding-ID
	excel_report_sheet.set_column(1, 1, 8)   # Fixed
	excel_report_sheet.set_column(2, 2, 20)  # Severity
	excel_report_sheet.set_column(3, 3, 20)  # Asset
	excel_report_sheet.set_column(4, 4, 60)  # Title

	# Findings
	row = 5
	col = 0 
	for counter,finding in enumerate(findings):
		# ID
		excel_report_sheet.write(row, col, display_finding_id(finding, counter), bold)

		# Fixed with color
		is_fixed = bool(finding.get("fixed", False))
		fixed_text = "Yes" if is_fixed else "No"
		fixed_fmt = fixed_yes_fmt if is_fixed else fixed_no_fmt
		excel_report_sheet.write(row, col + 1, fixed_text, fixed_fmt)

		# Severity with color
		sev = finding["cvss_severity"]
		sev_fmt = {
			"Critical": sev_crit_fmt,
			"High": sev_high_fmt,
			"Medium": sev_med_fmt,
			"Low": sev_low_fmt,
			"None": sev_none_fmt
		}.get(sev, None)
		sev_text = "{} ({})".format(sev, finding["cvss_score"])
		if sev_fmt:
			excel_report_sheet.write(row, col + 2, sev_text, sev_fmt)
		else:
			excel_report_sheet.write(row, col + 2, sev_text)

		# Asset and Title
		excel_report_sheet.write(row, col + 3, finding["asset"])
		excel_report_sheet.write(row, col + 4, finding["title"])
		row += 1

	excel_report.close()

def generate_word_report():
	"""Generate a Word document containing the sorted findings."""
	global config, findings, output_dir

	if not findings:
		process_findings()

	document = Document()

	active_findings = [f for f in findings if not f.get("fixed", False)]

	for idx, finding in enumerate(active_findings):
		display_id = finding.get("_display_id", display_finding_id(finding, idx))
		document.add_heading("{} {}".format(display_id, finding["title"]), level=1)

		cvss_score = finding["cvss_score"]
		if isinstance(cvss_score, (int, float)):
			cvss_score_text = "{:.1f}".format(cvss_score)
		else:
			cvss_score_text = str(cvss_score)

		document.add_paragraph("{}\t{} ({})".format(finding["cvss_vector"], finding["cvss_severity"], cvss_score_text))

		table_rows = [
			("Asset", finding["asset"]),
			("CWE", "CWE-{} ({})".format(finding["CWE-ID"], finding["CWE-Link"])),
			("Status", "Fixed" if finding.get("fixed", False) else "Open"),
			("Severity", "{} ({})".format(finding["cvss_severity"], cvss_score_text)),
			("CVSS Vector", finding["cvss_vector"])
		]

		table = document.add_table(rows=len(table_rows), cols=2)
		try:
			table.style = "Light List Accent 1"
		except KeyError:
			pass

		for row_idx, (label, value) in enumerate(table_rows):
			table.rows[row_idx].cells[0].text = label
			table.rows[row_idx].cells[1].text = str(value)

		document.add_paragraph("")
		add_markdown_content_to_doc(document, finding["description"].strip())

		if idx < len(active_findings) - 1:
			document.add_page_break()

	if not active_findings:
		document.add_paragraph("No active findings available.")

	output_path = os.path.join(output_dir, "findings.docx")
	document.save(output_path)
	print("Word report written to {}".format(output_path))


def generate_pdf_report(report_html, mode = "report", filename = "finding.md"):
	"""Generate PDF Report from HTML"""
	global boilerplate_dir, cover_location, output_dir
	# Generate PDF
	toc = {
		'xsl-style-sheet': boilerplate_dir + 'toc.xsl'
	}

	options = {
		'--header-html': boilerplate_dir + 'header.html',
		'--footer-html': boilerplate_dir + 'footer.html',
		#'footer-right': '[page] of [topage]',
		'footer-right': '[page]',
		'footer-font-name': 'avenir next',
		'footer-font-size': '10',
		'margin-bottom': '1.25cm', 
		'margin-top': '2.5cm',
		'header-spacing': '-5',
		'encoding': "UTF-8",
		'page-size': 'A4',
		"enable-local-file-access": None
	}

	css = boilerplate_dir + "report.css"

	print("Generating PDF...")
	if mode == "report":
		pdfkit.from_string(report_html, output_dir+'report.pdf', options=options, css=css, toc=toc, cover=cover_location, cover_first=True)
	elif mode == "findings":
		pdfkit.from_string(report_html, output_dir+filename, options=options, css=css)

def all():
	"""Generate all reports"""
	generate_report()
	generate_excel_report()


def print_findings():
	global config, findings, findings_dir, total_findings, critical_findings, high_findings, medium_findings, low_findings, none_findings

	print("Processed {} findings:".format(total_findings))
	print("Critical: {}".format(critical_findings))
	print("High: {}".format(high_findings))
	print("Medium: {}".format(medium_findings))
	print("Low: {}".format(low_findings))
	print("None: {}".format(none_findings))

	print("Findings:")
	for finding in findings:
		print("++++++++++++")
		print("Title: {}".format(finding["title"]))
		print("Asset: {}".format(finding["asset"]))
		print("Severity: {}".format(finding["cvss_severity"]))
		print("CVSS Score: {}".format(finding["cvss_score"]))
		print("")

def generate_findings_reports():
	"""Generate separate report files for all findings"""
	global config, findings, base_href

	for counter,finding in enumerate(findings):
		raw_finding_id = finding.get("_resolved_id", resolve_finding_id(finding, counter))
		display_id = finding.get("_display_id", display_finding_id(finding, counter))
		print("Generating report for finding {}...".format(display_id))
		finding_markdown_temp = finding_markdown(finding, display_id)
		finding_html = markdown.markdown(finding_markdown_temp, extensions=['fenced_code', 'codehilite', 'tables'])
		finding_html = wrap_html_document(finding_html, base_href)
		filename_id = raw_finding_id[1:] if raw_finding_id.startswith("#") else raw_finding_id
		generate_pdf_report(finding_html, mode = "findings", filename = "finding_{}.pdf".format(filename_id))

################################################

if __name__ == '__main__':
	init()

	# Parse arguments
	parser = argparse.ArgumentParser(description='Render a pentest report.')
	parser.add_argument('--all', default=False, action='store_true', help='Generate all reports from scratch.')
	parser.add_argument('--view_findings', default=False, action='store_true', help='Print all findings.')
	parser.add_argument('--findings_only', default=False, action='store_true', help='Generate separate report files for all findings.')
	parser.add_argument('--word', default=False, action='store_true', help='Export the sorted findings to a Word document.')
	if len(sys.argv) == 1:
		parser.print_help(sys.stderr)
		sys.exit(1)
	args = parser.parse_args()

	if args.all:
		all()
	
	if args.view_findings:
		process_findings()
		print_findings()

	if args.findings_only:
		process_findings()
		generate_findings_reports()

	if args.word:
		generate_word_report()
